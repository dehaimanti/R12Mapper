# rtf_generator.py
import os
from docx import Document
from docx.shared import Inches

def generate_rtf_template(xml_data_path, output_path):
    """
    Creates a simple BI Publisher-compatible RTF using python-docx.
    Assumes standard PO layout. Does not embed data — only generates layout.
    """
    doc = Document()
    doc.add_heading('Purchase Order Template', 0)

    doc.add_paragraph('Number: <?Header/Number?>')
    doc.add_paragraph('Type: <?Header/Type?>')
    doc.add_paragraph('Order / Revision: <?Header/OrderRevision?>')
    doc.add_paragraph('Approved Date: <?Header/ApprovedDate?>')
    doc.add_paragraph('Created By: <?Header/CreatedBy?>')
    doc.add_paragraph('Buyer: <?Header/Buyer?>')
    doc.add_paragraph('Buyer Email: <?Header/BuyerEmail?>')

    doc.add_paragraph('\nLine Items:')
    doc.add_paragraph('<?for-each:Lines/Line?>')

    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Part Number'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Quantity'
    hdr_cells[3].text = 'UOM'
    hdr_cells[4].text = 'Unit Price'
    hdr_cells[5].text = 'Total'

    doc.add_paragraph('<?end for-each?>')

    doc.save(output_path)
    return output_path
